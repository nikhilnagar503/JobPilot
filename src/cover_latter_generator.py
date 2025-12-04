import os
import PyPDF2
from docx import Document
from src.nlp_processing import extract_skills_from_description
from datetime import datetime
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders
from dotenv import load_dotenv
import json
from langchain_groq import ChatGroq
from langchain.prompts import PromptTemplate

load_dotenv()

# Load the .env file
load_dotenv(dotenv_path=r"C:\Users\nagar\Desktop\my_project\job_automate\JobPilot\gmail.env")

llm = ChatGroq(model="llama3-8b-8192", groq_api_key=os.getenv("GROQ_API_KEY"))


def generate_cover_letter(job_title, company, job_desc, cv_file_path):
    # Extract skills from job description
    skills = extract_skills_from_description(job_desc)
    
    # Extract relevant experience from CV
    experience = extract_experience_from_cv(cv_file_path)

    # Extract name and contact info from the CV
    name, contact_info = extract_name_and_contact_from_cv(cv_file_path)

    # Create the cover letter template
    cover_letter = f"""
    Dear Hiring Manager,

    I am excited to apply for the {job_title} position at {company}. With a strong background in {', '.join(skills)}, 
    I am eager to contribute my expertise to your team.

    Currently, I am a {experience}. I have successfully contributed to optimizing experimental workflows, improving predictive accuracy, and co-authoring research papers. Furthermore, I have disseminated AI/ML research publications to a community of over 200 data scientists, helping foster collaboration and drive advancements in the field.

    I am excited about the opportunity to apply my technical skills, analytical expertise, and academic background to the {job_title} role at {company}, where I can leverage my experience to help drive positive change in the advertising industry. The opportunity to work on flexible hours and contribute to a team that ensures responsible advertising resonates with my professional values and long-term career goals.

    I look forward to the opportunity to further discuss how my background and skills can contribute to the continued success of {company}. Thank you for your consideration.

    Sincerely,  
    {name}  
    {contact_info}
    """
    
    return cover_letter

# Helper functions for extracting information from CV
def extract_experience_from_cv(cv_file_path):
    experience = ""

    if cv_file_path.lower().endswith('.pdf'):
        with open(cv_file_path, 'rb') as cv_file:
            reader = PyPDF2.PdfReader(cv_file)
            text = ""
            for page in reader.pages:
                text += page.extract_text()
        experience = extract_experience_with_langchain(text)

    elif cv_file_path.lower().endswith('.docx'):
        doc = Document(cv_file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        experience = extract_experience_with_langchain(text)

    return experience

def extract_experience_with_langchain(cv_text):
    prompt = f"""
    Extract ONLY the EXPERIENCE section from this CV text.
    Return short summary of work experience, roles, or responsibilities.

    If no experience is found, return an empty string.

    Return ONLY this format:

    Experience: <experience_text_here>

    CV Text:
    {cv_text}
    """

    response = llm.predict(prompt)

    # default empty
    experience = ""

    # parse line
    for line in response.split("\n"):
        if line.lower().startswith("experience"):
            experience = line.split(":", 1)[1].strip()

    return experience


def extract_name_and_contact_from_cv(cv_file_path):
    name = ""
    contact_info = ""

    if cv_file_path.lower().endswith('.pdf'):
        with open(cv_file_path, 'rb') as cv_file:
            reader = PyPDF2.PdfReader(cv_file)
            text = ""
            for page in reader.pages:
                text += page.extract_text()
        name, contact_info = extract_name_and_contact_with_langchain(text)

    elif cv_file_path.lower().endswith('.docx'):
        doc = Document(cv_file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        name, contact_info = extract_name_and_contact_with_langchain(text)

    return name, contact_info

def extract_name_and_contact_with_langchain(cv_text):
    prompt = f"""
    Extract ONLY the NAME and CONTACT details (Email and Phone) from the CV text.

    Return ONLY this format:

    Name: <name>
    Contact: <email_or_phone>

    If any detail is missing, leave it blank.

    CV Text:
    {cv_text}
    """

    response = llm.predict(prompt)

    # DEFAULT EMPTY VALUES
    name = ""
    contact = ""

    # Parse line by line
    for line in response.split("\n"):
        if line.lower().startswith("name"):
            name = line.split(":", 1)[1].strip()
        if line.lower().startswith("contact"):
            contact = line.split(":", 1)[1].strip()

    return name, contact

# Save the CV and Cover Letter to Files
def save_to_files(cv_file, cover_letter, name):
    output_dir = "generated_documents"
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    cover_letter_filename = f"Cover_letter_{name}.txt"
    cv_filename = f"CV_{name}.txt"
    
    with open(os.path.join(output_dir, cover_letter_filename), "w") as f:
        f.write(cover_letter)
    
    with open(os.path.join(output_dir, cv_filename), "w") as f:
        f.write("CV content goes here...")  # This should ideally be CV content extracted from the file

    return os.path.join(output_dir, cover_letter_filename), os.path.join(output_dir, cv_filename)

# Function to send email with CV and Cover Letter as attachments
def send_email(subject, body, recipient, cv_path, cover_letter_path):
    msg = MIMEMultipart()
    msg['From'] = os.getenv('EMAIL_USER')
    msg['To'] = recipient
    msg['Subject'] = subject

    msg.attach(MIMEText(body, 'plain'))

    try:
        with open(cv_path, 'rb') as file:
            part = MIMEBase('application', 'octet-stream')
            part.set_payload(file.read())
            encoders.encode_base64(part)
            part.add_header('Content-Disposition', f'attachment; filename={os.path.basename(cv_path)}')
            msg.attach(part)
    except Exception as e:
        print(f"Error attaching CV: {e}")

    try:
        with open(cover_letter_path, 'rb') as file:
            part = MIMEBase('application', 'octet-stream')
            part.set_payload(file.read())
            encoders.encode_base64(part)
            part.add_header('Content-Disposition', f'attachment; filename={os.path.basename(cover_letter_path)}')
            msg.attach(part)
    except Exception as e:
        print(f"Error attaching cover letter: {e}")

    try:
        server = smtplib.SMTP('smtp.gmail.com', 587)
        server.starttls()
        server.login(os.getenv('EMAIL_USER'), os.getenv('EMAIL_PASSWORD'))
        server.sendmail(msg['From'], recipient, msg.as_string())
        server.quit()
        print("Email sent successfully!")
    except Exception as e:
        print(f"Error sending email: {e}")
